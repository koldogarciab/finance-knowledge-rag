from pathlib import Path
import json

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUEPRINT_PATH = Path("data/blueprint/corpus_blueprint.json")
OUTPUT_PATH = Path(
    "data/raw/docx/finance_policies_and_procedures.docx"
)

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "E7E6E6"
DARK_GREY = RGBColor(89, 89, 89)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))

    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)

    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    row_properties.append(table_header)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.extend([begin, instruction, separate, text, end])


def format_aud(value) -> str:
    if value is None:
        return "No upper limit"

    return f"AUD {value:,.0f}"


def style_table_header(row) -> None:
    set_repeat_table_header(row)

    for cell in row.cells:
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)


def add_bullet(document, text: str) -> None:
    paragraph = document.add_paragraph(
        text,
        style="List Bullet"
    )
    paragraph.paragraph_format.space_after = Pt(3)


with BLUEPRINT_PATH.open("r", encoding="utf-8") as file:
    data = json.load(file)

OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

people_by_id = {
    person["person_id"]: person
    for person in data["people"]
}

finance_director = people_by_id["PER_002"]
financial_controller = people_by_id["PER_004"]
head_of_fpa = people_by_id["PER_003"]

document = Document()

section = document.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.70)
section.left_margin = Inches(0.80)
section.right_margin = Inches(0.80)

document.core_properties.title = "Finance Policies and Procedures"
document.core_properties.subject = (
    "Synthetic finance policies for Harbour Retail Group"
)
document.core_properties.author = "Harbour Retail Group"
document.core_properties.keywords = (
    "finance, policy, controls, approvals, month-end, synthetic"
)

styles = document.styles

normal_style = styles["Normal"]
normal_style.font.name = "Arial"
normal_style.font.size = Pt(10)
normal_style.paragraph_format.space_after = Pt(6)

for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2"]:
    styles[style_name].font.name = "Arial"

styles["Title"].font.size = Pt(26)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor(31, 78, 121)

styles["Subtitle"].font.size = Pt(14)
styles["Subtitle"].font.color.rgb = DARK_GREY

styles["Heading 1"].font.size = Pt(16)
styles["Heading 1"].font.bold = True
styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
styles["Heading 1"].paragraph_format.space_before = Pt(12)
styles["Heading 1"].paragraph_format.space_after = Pt(6)
styles["Heading 1"].paragraph_format.keep_with_next = True

styles["Heading 2"].font.size = Pt(12)
styles["Heading 2"].font.bold = True
styles["Heading 2"].font.color.rgb = RGBColor(68, 68, 68)
styles["Heading 2"].paragraph_format.keep_with_next = True

# Header and footer
header = section.header.paragraphs[0]
header.text = (
    "Harbour Retail Group | Finance Policies and Procedures"
)
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

for run in header.runs:
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = DARK_GREY

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

footer_run = footer.add_run(
    "Synthetic finance corpus | Page "
)
footer_run.font.name = "Arial"
footer_run.font.size = Pt(8)
footer_run.font.color.rgb = DARK_GREY

add_page_number(footer)

# ------------------------------------------------------------------
# Cover page
# ------------------------------------------------------------------

document.add_paragraph("")
document.add_paragraph("")
document.add_paragraph("HARBOUR RETAIL GROUP", style="Subtitle")

title = document.add_paragraph(
    "Finance Policies and Procedures",
    style="Title"
)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = document.add_paragraph(
    f"{data['financial_year']['label']} | Version 1.0",
    style="Subtitle"
)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph("")
document.add_paragraph("")

cover_table = document.add_table(rows=5, cols=2)
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_table.style = "Table Grid"

cover_values = [
    ("Document ID", "DOC_DOCX_001"),
    ("Policy owner", finance_director["name"]),
    ("Effective date", "1 July 2025"),
    ("Review cycle", "Annual"),
    ("Classification", "Internal - synthetic training data")
]

for row, (label, value) in zip(cover_table.rows, cover_values):
    row.cells[0].text = label
    row.cells[1].text = value
    set_cell_shading(row.cells[0], LIGHT_BLUE)

    for run in row.cells[0].paragraphs[0].runs:
        run.font.bold = True

document.add_paragraph("")

notice = document.add_paragraph()
notice.alignment = WD_ALIGN_PARAGRAPH.CENTER

notice_run = notice.add_run(
    "This document contains synthetic business information created "
    "for a reproducible RAG evaluation project."
)
notice_run.font.italic = True
notice_run.font.color.rgb = DARK_GREY
notice_run.font.size = Pt(9)

document.add_page_break()

# ------------------------------------------------------------------
# Document overview
# ------------------------------------------------------------------

document.add_heading("1. Document overview", level=1)

document.add_heading("1.1 Purpose", level=2)
document.add_paragraph(
    "This document defines the principal finance controls, approval "
    "requirements, accounting procedures and forecasting responsibilities "
    "used by Harbour Retail Group."
)

document.add_heading("1.2 Scope", level=2)
document.add_paragraph(
    "The policies apply to all Harbour Retail Group departments, stores, "
    "corporate functions and employees involved in expenditure, financial "
    "reporting, procurement, forecasting or record retention."
)

document.add_heading("1.3 Responsibilities", level=2)

responsibility_table = document.add_table(rows=1, cols=2)
responsibility_table.style = "Table Grid"
responsibility_table.alignment = WD_TABLE_ALIGNMENT.CENTER

responsibility_table.rows[0].cells[0].text = "Role"
responsibility_table.rows[0].cells[1].text = "Primary responsibility"
style_table_header(responsibility_table.rows[0])

responsibilities = [
    (
        f"{finance_director['title']} - {finance_director['name']}",
        "Owns the finance policy framework and approves material "
        "operating and capital expenditure."
    ),
    (
        f"{financial_controller['title']} - "
        f"{financial_controller['name']}",
        "Owns month-end close, journals, accruals, prepayments, "
        "reconciliations and finance document retention."
    ),
    (
        f"{head_of_fpa['title']} - {head_of_fpa['name']}",
        "Owns departmental forecasting, consolidation, challenge "
        "and executive forecast reporting."
    ),
    (
        "Department Directors",
        "Ensure expenditure is authorised, forecasts are complete "
        "and supporting documentation is retained."
    )
]

for role, responsibility in responsibilities:
    cells = responsibility_table.add_row().cells
    cells[0].text = role
    cells[1].text = responsibility

document.add_heading("1.4 Policy index", level=2)

index_table = document.add_table(rows=1, cols=3)
index_table.style = "Table Grid"
index_table.alignment = WD_TABLE_ALIGNMENT.CENTER

index_table.rows[0].cells[0].text = "Policy ID"
index_table.rows[0].cells[1].text = "Section"
index_table.rows[0].cells[2].text = "Policy title"
style_table_header(index_table.rows[0])

for policy in data["finance_policies"]:
    cells = index_table.add_row().cells
    cells[0].text = policy["policy_id"]
    cells[1].text = policy["section"]
    cells[2].text = policy["title"]

# ------------------------------------------------------------------
# Individual policies
# ------------------------------------------------------------------

for number, policy in enumerate(
    data["finance_policies"],
    start=2
):
    document.add_page_break()

    document.add_heading(
        f"{number}. {policy['title']}",
        level=1
    )

    owner = people_by_id[policy["owner_person_id"]]

    metadata_table = document.add_table(rows=3, cols=2)
    metadata_table.style = "Table Grid"
    metadata_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    metadata = [
        ("Policy ID", policy["policy_id"]),
        ("Policy section", policy["section"]),
        (
            "Policy owner",
            f"{owner['name']} - {owner['title']}"
        )
    ]

    for row, (label, value) in zip(metadata_table.rows, metadata):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_BLUE)

        for run in row.cells[0].paragraphs[0].runs:
            run.font.bold = True

    document.add_paragraph("")

    if policy.get("rules"):
        document.add_heading("Policy requirements", level=2)

        for rule in policy["rules"]:
            add_bullet(document, rule)

    if policy.get("approval_levels"):
        document.add_heading("Approval levels", level=2)

        approval_table = document.add_table(rows=1, cols=3)
        approval_table.style = "Table Grid"
        approval_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = [
            "Minimum amount",
            "Maximum amount",
            "Required approvers"
        ]

        for index, header_text in enumerate(headers):
            approval_table.rows[0].cells[index].text = header_text

        style_table_header(approval_table.rows[0])

        for level in policy["approval_levels"]:
            cells = approval_table.add_row().cells

            cells[0].text = format_aud(level["minimum_aud"])
            cells[1].text = format_aud(level["maximum_aud"])
            cells[2].text = "; ".join(
                level["required_approvers"]
            )

            for cell in cells:
                cell.vertical_alignment = (
                    WD_CELL_VERTICAL_ALIGNMENT.CENTER
                )

    if policy.get("additional_rules"):
        document.add_heading(
            "Additional requirements",
            level=2
        )

        for rule in policy["additional_rules"]:
            add_bullet(document, rule)

# ------------------------------------------------------------------
# Appendix
# ------------------------------------------------------------------

document.add_page_break()
document.add_heading(
    "12. Appendix - key thresholds",
    level=1
)

threshold_table = document.add_table(rows=1, cols=3)
threshold_table.style = "Table Grid"
threshold_table.alignment = WD_TABLE_ALIGNMENT.CENTER

threshold_headers = [
    "Control",
    "Threshold",
    "Requirement"
]

for index, header_text in enumerate(threshold_headers):
    threshold_table.rows[0].cells[index].text = header_text

style_table_header(threshold_table.rows[0])

thresholds = [
    (
        "Accrual",
        "AUD 5,000",
        "Minimum individual accrual threshold"
    ),
    (
        "Prepayment",
        "AUD 12,000",
        "Must also provide benefit for more than three months"
    ),
    (
        "Purchase order",
        "Above AUD 5,000",
        "Must be approved before committing expenditure"
    ),
    (
        "Manual journal",
        "Above AUD 100,000",
        "Requires Financial Controller approval"
    ),
    (
        "Capex business case",
        "Above AUD 50,000",
        "Documented business case required"
    ),
    (
        "Board capex approval",
        "Above AUD 250,000",
        "Board of Directors approval required"
    ),
    (
        "Forecast commentary",
        "Above AUD 100,000",
        "Written commentary required"
    ),
    (
        "Record retention",
        "Seven years",
        "Accounting records and supporting documentation"
    )
]

for control, threshold, requirement in thresholds:
    cells = threshold_table.add_row().cells
    cells[0].text = control
    cells[1].text = threshold
    cells[2].text = requirement

document.add_paragraph("")

final_note = document.add_paragraph()
final_note.alignment = WD_ALIGN_PARAGRAPH.CENTER

final_run = final_note.add_run(
    "End of Finance Policies and Procedures"
)
final_run.font.bold = True
final_run.font.color.rgb = RGBColor(31, 78, 121)

document.save(OUTPUT_PATH)

print("Finance policies DOCX generated successfully")
print(f"Policies: {len(data['finance_policies'])}")
print(f"Tables: {len(document.tables)}")
print(f"Created: {OUTPUT_PATH}")
print(f"File size: {OUTPUT_PATH.stat().st_size:,} bytes")
