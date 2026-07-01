from pathlib import Path
from zipfile import ZipFile
import json
import sys

from docx import Document


BLUEPRINT_PATH = Path("data/blueprint/corpus_blueprint.json")
DOCX_PATH = Path(
    "data/raw/docx/finance_policies_and_procedures.docx"
)


def fail(message: str) -> None:
    print(f"[FAIL] {message}")
    sys.exit(1)


if not DOCX_PATH.exists():
    fail(f"DOCX file is missing: {DOCX_PATH}")

if DOCX_PATH.stat().st_size == 0:
    fail("DOCX file is empty")


with BLUEPRINT_PATH.open("r", encoding="utf-8") as file:
    blueprint = json.load(file)

try:
    document = Document(DOCX_PATH)
except Exception as error:
    fail(f"DOCX cannot be opened: {error}")


paragraph_text = "\n".join(
    paragraph.text
    for paragraph in document.paragraphs
)

table_text = "\n".join(
    cell.text
    for table in document.tables
    for row in table.rows
    for cell in row.cells
)

full_text = f"{paragraph_text}\n{table_text}"


# -------------------------------------------------------------------
# General document structure
# -------------------------------------------------------------------

required_document_values = [
    "HARBOUR RETAIL GROUP",
    "Finance Policies and Procedures",
    "FY2025/26 | Version 1.0",
    "DOC_DOCX_001",
    "Internal - synthetic training data",
    "1. Document overview",
    "1.1 Purpose",
    "1.2 Scope",
    "1.3 Responsibilities",
    "1.4 Policy index",
    "12. Appendix - key thresholds",
    "End of Finance Policies and Procedures"
]

for value in required_document_values:
    if value not in full_text:
        fail(f"Required document text is missing: {value}")


if len(document.tables) != 16:
    fail(
        f"Expected 16 tables, found {len(document.tables)}"
    )


# -------------------------------------------------------------------
# Validate every policy
# -------------------------------------------------------------------

people_by_id = {
    person["person_id"]: person
    for person in blueprint["people"]
}

policies = blueprint["finance_policies"]

if len(policies) != 10:
    fail(
        f"Blueprint should contain 10 policies, "
        f"found {len(policies)}"
    )

for policy in policies:
    policy_id = policy["policy_id"]
    owner = people_by_id[policy["owner_person_id"]]

    required_values = [
        policy_id,
        policy["section"],
        policy["title"],
        owner["name"],
        owner["title"]
    ]

    for value in required_values:
        if value not in full_text:
            fail(
                f"{policy_id}: required value is missing: "
                f"{value}"
            )

    for rule in policy.get("rules", []):
        if rule not in full_text:
            fail(
                f"{policy_id}: policy rule is missing: "
                f"{rule}"
            )

    for rule in policy.get("additional_rules", []):
        if rule not in full_text:
            fail(
                f"{policy_id}: additional rule is missing: "
                f"{rule}"
            )

    for level in policy.get("approval_levels", []):
        for approver in level["required_approvers"]:
            if approver not in full_text:
                fail(
                    f"{policy_id}: approver is missing: "
                    f"{approver}"
                )


# -------------------------------------------------------------------
# Validate key thresholds
# -------------------------------------------------------------------

required_thresholds = {
    "Accrual": "AUD 5,000",
    "Prepayment": "AUD 12,000",
    "Purchase order": "Above AUD 5,000",
    "Manual journal": "Above AUD 100,000",
    "Capex business case": "Above AUD 50,000",
    "Board capex approval": "Above AUD 250,000",
    "Forecast commentary": "Above AUD 100,000",
    "Record retention": "Seven years"
}

for control, threshold in required_thresholds.items():
    if control not in full_text:
        fail(f"Threshold control is missing: {control}")

    if threshold not in full_text:
        fail(
            f"Threshold value is missing for {control}: "
            f"{threshold}"
        )


# -------------------------------------------------------------------
# Validate header, footer and page-number field
# -------------------------------------------------------------------

header_text = "\n".join(
    paragraph.text
    for section in document.sections
    for paragraph in section.header.paragraphs
)

footer_text = "\n".join(
    paragraph.text
    for section in document.sections
    for paragraph in section.footer.paragraphs
)

if (
    "Harbour Retail Group | Finance Policies and Procedures"
    not in header_text
):
    fail("Expected document header is missing")

if "Synthetic finance corpus | Page" not in footer_text:
    fail("Expected document footer is missing")

with ZipFile(DOCX_PATH, "r") as archive:
    file_names = set(archive.namelist())

    required_internal_files = {
        "[Content_Types].xml",
        "word/document.xml",
        "word/styles.xml",
        "word/header1.xml",
        "word/footer1.xml"
    }

    missing_files = required_internal_files - file_names

    if missing_files:
        fail(
            "DOCX package is missing internal files: "
            f"{sorted(missing_files)}"
        )

    footer_xml = archive.read(
        "word/footer1.xml"
    ).decode("utf-8")

    if " PAGE " not in footer_xml:
        fail("Footer does not contain a PAGE field")


print("[PASS] Finance policies DOCX opened successfully")
print(f"[PASS] File size: {DOCX_PATH.stat().st_size:,} bytes")
print(f"[PASS] Tables: {len(document.tables)}")
print(f"[PASS] Policies: {len(policies)}")
print("[PASS] Document overview and policy index are complete")
print("[PASS] Every policy matches the blueprint")
print("[PASS] Approval levels and approvers are complete")
print("[PASS] Appendix thresholds are complete")
print("[PASS] Header, footer and page-number field are present")
print("[PASS] Finance policies DOCX is fully consistent")
